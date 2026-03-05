from __future__ import annotations

import argparse
import json
import sys
import time
import uuid
from pathlib import Path

import httpx


DEFAULT_ALLOWED_EXTS = ".txt,.md,.py,.json,.csv,.pdf,.docx"
FINAL_STATUSES = {"SUCCEEDED", "FAILED", "DEAD_LETTER"}
TEXT_FILE_EXTS = {".txt", ".md", ".py", ".json", ".csv"}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Scan a local folder, send files as inline_documents to ingestion, and poll job status.",
    )
    parser.add_argument("--folder", required=True, help="Local folder path to scan recursively.")
    parser.add_argument("--ingestion-url", default="http://127.0.0.1:8082", help="Ingestion API base URL.")
    parser.add_argument("--tenant-id", default="tenant-a", help="tenant_id for ingestion job.")
    parser.add_argument("--app-id", default="app-a", help="app_id for ingestion job.")
    parser.add_argument("--source-id", default="local-folder-001", help="source.source_id for ingestion job.")
    parser.add_argument("--source-type", default="upload", help="source.source_type for ingestion job.")
    parser.add_argument(
        "--allowed-exts",
        default=DEFAULT_ALLOWED_EXTS,
        help=f"Comma-separated extensions to include (default: {DEFAULT_ALLOWED_EXTS}).",
    )
    parser.add_argument("--preferred-embedding-model", default="bge-m3", help="options.preferred_embedding_model.")
    parser.add_argument("--create-timeout", type=float, default=60.0, help="HTTP timeout for job creation request.")
    parser.add_argument("--status-timeout", type=float, default=30.0, help="HTTP timeout for status polling request.")
    parser.add_argument("--poll-interval", type=float, default=2.0, help="Seconds between status polls.")
    return parser.parse_args()


def normalize_extensions(raw: str) -> set[str]:
    exts: set[str] = set()
    for item in raw.split(","):
        ext = item.strip().lower()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = f".{ext}"
        exts.add(ext)
    return exts


def infer_file_type(ext: str) -> str:
    ext = ext.lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "word"
    return "text"


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("missing dependency: pypdf") from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())
    return "\n".join(parts).strip()


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("missing dependency: python-docx") from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in TEXT_FILE_EXTS:
        return path.read_text(encoding="utf-8", errors="ignore").strip()
    if ext == ".pdf":
        return extract_pdf_text(path)
    if ext == ".docx":
        return extract_docx_text(path)
    raise RuntimeError(f"unsupported extension: {ext}")


def collect_documents(folder: Path, allowed_exts: set[str]) -> list[dict]:
    docs: list[dict] = []
    for path in folder.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in allowed_exts:
            continue
        try:
            content = extract_text(path)
        except Exception as exc:
            print(f"[warn] skip file parse failed: {path} ({exc})", file=sys.stderr)
            continue
        if not content:
            continue
        rel = path.relative_to(folder).as_posix()
        docs.append(
            {
                "doc_id": rel.replace("/", "__"),
                "title": path.stem,
                "file_name": rel,
                "file_type": infer_file_type(path.suffix),
                "content": content,
            }
        )
    return docs


def main() -> int:
    args = parse_args()
    folder = Path(args.folder).resolve()
    if not folder.is_dir():
        print(f"[error] folder not found: {folder}", file=sys.stderr)
        return 2

    allowed_exts = normalize_extensions(args.allowed_exts)
    docs = collect_documents(folder, allowed_exts)
    if not docs:
        print(f"[error] no eligible files found under: {folder}", file=sys.stderr)
        return 2

    print(f"[info] collected docs={len(docs)} from {folder}")
    payload = {
        "trace_id": f"t-{uuid.uuid4().hex[:8]}",
        "request_id": f"r-{uuid.uuid4().hex[:8]}",
        "tenant_id": args.tenant_id,
        "app_id": args.app_id,
        "source": {"source_id": args.source_id, "source_type": args.source_type},
        "options": {
            "preferred_embedding_model": args.preferred_embedding_model,
            "inline_documents": docs,
        },
    }

    create_url = f"{args.ingestion_url.rstrip('/')}/api/v0.1/ingestion/jobs"
    with httpx.Client() as client:
        response = client.post(create_url, json=payload, timeout=args.create_timeout)
        response.raise_for_status()
        created = response.json()
        job_id = created["job_id"]
        print(f"[info] created job_id={job_id}")

        status_url = f"{args.ingestion_url.rstrip('/')}/api/v0.1/ingestion/jobs/{job_id}"
        while True:
            status_resp = client.get(status_url, timeout=args.status_timeout)
            status_resp.raise_for_status()
            status = status_resp.json()
            print(
                f"[status] {status['status']} stage={status.get('current_stage')} "
                f"stats={json.dumps(status.get('stats', {}), ensure_ascii=False)}"
            )
            if status["status"] in FINAL_STATUSES:
                print("[result]", json.dumps(status, ensure_ascii=False))
                break
            time.sleep(args.poll_interval)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
